


import os
from crewai import Crew, Process
from dotenv import load_dotenv

from agents import (
    requirement_analyst_Agent,
    user_story_agent,
    pdd_agent,
    sdd_agent,
    component_mapper_agent,
    risk_analysis_agent,
    flowchart_agent,
    document_reviewer_agent,
    document_updater_agent,
    final_writer_agent,
    word_doc_summary_agent,
    evaluation_agent,
    word_doc_full_agent
)

from tasks import (
    requirement_analyst_task,
    requirement_analyst_review_task,
    requirement_update_task,
    user_story_task,
    user_story_review_task,
    user_story_update_task,
    pdd_task,
    pdd_review_task,
    pdd_update_task,
    sdd_task,
    sdd_review_task,
    sdd_update_task,
    component_map_task,
    risk_analysis_task,
    flowchart_task,
    final_writer_task,
    word_doc_summary_task,    # ✅ Uses only final_writer output
    word_doc_full_task,
    evaluation_task        # ✅ Uses full updated context
)


import re
from docx import Document

load_dotenv()

print("📘 Welcome to the Agentic AI Requirement Understanding System!")

# === Step 1: Clean output directory ===
output_folder = "outputs"
os.makedirs(output_folder, exist_ok=True)
for file in os.listdir(output_folder):
    file_path = os.path.join(output_folder, file)
    if os.path.isfile(file_path):
        os.remove(file_path)


# # === Step 2: Define helper functions for DOCX and PDF ===

def generate_word_doc(input_path, output_path, title):
    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    doc.add_heading(title, 0)

    for line in lines:
        stripped = line.strip()

        if not stripped:
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("* ") or stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped.startswith("**") and stripped.endswith("**:"):
            doc.add_heading(stripped.strip("*:"), level=3)
        else:
            doc.add_paragraph(stripped)

    doc.save(output_path)
    print(f"✅ DOCX generated: {output_path}")




# === Step 3: Initialize Crew ===

crew = Crew(
    agents=[
        requirement_analyst_Agent,
        user_story_agent,
        pdd_agent,
        component_mapper_agent,
        sdd_agent,
        risk_analysis_agent,
        flowchart_agent,
        document_reviewer_agent,
        document_updater_agent,
        final_writer_agent,
        evaluation_agent,
        word_doc_summary_agent,
        
        word_doc_full_agent
    ],

    tasks=[
        requirement_analyst_task,
        requirement_analyst_review_task,
        requirement_update_task,
        user_story_task,
        user_story_review_task,
        user_story_update_task,
        pdd_task,
        pdd_review_task,
        pdd_update_task,
        component_map_task,
        sdd_task,
        sdd_review_task,
        sdd_update_task,
        risk_analysis_task,
        flowchart_task,
        final_writer_task,
        evaluation_task,     # ✅ Uses full updated context
        word_doc_summary_task,    # ✅ Uses only final_writer output
        word_doc_full_task        # ✅ Uses full updated context
    ],
    process=Process.sequential,
    verbose=True
)


# === Step 4: Run Crew and Generate Output Files ===

print("\n🚀 [STEP 1] Executing all tasks...")

try:
    result = crew.kickoff(inputs={
        'topic': "Design an AI-powered Document Verification System for a financial institution that automates "
        "the verification of KYC documents (such as Aadhar, PAN, Passport, Driving License) using OCR and computer"
        " vision techniques. The system should extract information from uploaded documents, validate them against "
        "the user’s entered data, detect potential forgeries, and provide a confidence score. It should integrate "
        "with the bank’s onboarding workflow and support a secure API for external integration."

    })

    print("\n✅ [STEP 2] All tasks completed.")







    # === Step 6: Generate documents from .txt ===

    generate_word_doc("outputs/docs_full.txt", "outputs/09a_final_full.docx", " Full Technical Report")
    generate_word_doc("outputs/docs_summary.txt", "outputs/09b_summary.docx", " Summarized Report")


    print("\n🎉 All documents successfully created!")

except Exception as e:
    print(f"❌ Error during execution: {e}")
    exit(1)







